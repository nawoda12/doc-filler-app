import streamlit as st
from docx import Document
from docx.shared import RGBColor, Inches
import re
from datetime import datetime

def fill_template(template_path, replacements, logo_path=None):
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph, replacements):
        original_text = paragraph.text
        new_text = original_text

        for placeholder, value in replacements.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)

        new_text = re.sub(r'<[^>]*>', '', new_text)
        new_text = re.sub(r'\[[^\]]*\]', '', new_text)

        if "Total Estimated Hours" in new_text and "xx Hours" in new_text and "Hours" in replacements:
            new_text = new_text.replace("xx Hours", replacements["Hours"])

        if any(key in original_text for key in replacements.keys()) or original_text.strip() != "":
            paragraph.text = new_text
            for run in paragraph.runs:
                if "THE MASTER AGREEMENT AND" not in run.text:
                    run.font.color.rgb = RGBColor(0, 0, 0)

    def replace_text_in_cell(cell, replacements):
        for paragraph in cell.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, replacements)

    # Replace text in text boxes (shapes)
    for shape in doc.inline_shapes:
        if shape.type == 3:  # Text box
            try:
                text_box = shape._inline.graphic.graphicData.txbxContent
                for paragraph in text_box.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
            except Exception:
                continue

    # Replace the first image (logo) with the uploaded one
    if logo_path:
        for rel in doc.part._rels:
            rel_obj = doc.part._rels[rel]
            if "image" in rel_obj.reltype:
                rel_obj._target = logo_path
                break

    return doc

def extract_replacements(email_content):
    replacements = {
        "<Name>": "Nawoda Sathsara"
    }
    lines = email_content.split('\n')
    for line in lines:
        match = re.match(r'(.+?):\s*(.+)', line)
        if match:
            key = match.group(1).strip()
            value = match.group(2).strip()
            replacements[f"<{key}>"] = value
            replacements[f"[{key}]"] = value
            replacements[key] = value
    replacements["_____________"] = replacements.get("Date", "DATE")
    replacements["_______________________________________________"] = replacements.get("Customer", "XYZ Co.")
    replacements["______________________________________________ _____________"] = replacements.get("Customer Address", "330, Flatbush Avenue, Brooklyn, New York 11238, USA")
    replacements["__________________"] = replacements.get("Contract Execution Date", "CONTRACT EXECUTION DATE")
    replacements["[year]"] = str(datetime.now().year)
    return replacements

def get_document_preview(doc):
    preview = ""
    for para in doc.paragraphs[:10]:  # Preview first 10 paragraphs
        preview += para.text + "\n"
    return preview

st.title("📄 Statement of Work Auto-Filler")

uploaded_file = st.file_uploader("Upload your Word template (.docx)", type="docx")
logo_file = st.file_uploader("Upload your logo (.png, .jpg)", type=["png", "jpg"])
email_content = st.text_area("Paste the client email content here")

if uploaded_file and email_content:
    replacements = extract_replacements(email_content)
    logo_path = None
    if logo_file:
        logo_path = "uploaded_logo." + logo_file.name.split('.')[-1]
        with open(logo_path, "wb") as f:
            f.write(logo_file.getbuffer())

    filled_doc = fill_template(uploaded_file, replacements, logo_path)


    # Save and offer download
    filled_doc_path = "Filled_Statement_of_Work.docx"
    filled_doc.save(filled_doc_path)

    with open(filled_doc_path, "rb") as file:
        st.download_button(
            label="📥 Download Filled Document",
            data=file,
            file_name="Filled_Statement_of_Work.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
